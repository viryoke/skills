#!/usr/bin/env python3
"""将 Markdown 文档转换为 Docx 格式，用于归档阶段（Phase 5）

核心特性：
- Markdown 保留图表源码（PlantUML/DOT/Mermaid 代码块）
- Docx 转换时自动渲染图表为 PNG 并嵌入
- 使用内置打包工具（plantuml.jar / graphviz），无需联网
- python-docx 替代 pandoc，自动 pip install

用法:
  单文件: python3 convert_docx.py <file.md> [docx_template]
  批量:   python3 convert_docx.py <output_dir> [docx_template]
例如: python3 convert_docx.py output references/docx_template.docx
"""

import os
import sys
import glob
import subprocess
import re
import shutil
import platform
import importlib

# ─── 依赖自动安装 ───

def ensure_dep(package_name, import_name=None):
    """确保 Python 依赖已安装，缺失时自动 pip install"""
    import_name = import_name or package_name
    try:
        importlib.import_module(import_name)
        return True
    except ImportError:
        print("[自动安装] " + package_name + "...")
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet", package_name],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            print("[错误] 安装 " + package_name + " 失败:")
            print("  " + result.stderr.strip())
            return False
        print("[已安装] " + package_name)
        importlib.invalidate_caches()
        importlib.import_module(import_name)
        return True


if not ensure_dep("python-docx", "docx"):
    print("无法继续，请手动执行: pip install python-docx")
    sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))


# ─── 图表渲染器 ───

class DiagramRenderer:
    """渲染 PlantUML / DOT / Mermaid 为 PNG，使用内置打包工具"""

    def __init__(self, output_dir):
        self.output_dir = output_dir
        self.diagrams_dir = os.path.join(output_dir, "diagrams")
        os.makedirs(self.diagrams_dir, exist_ok=True)
        self.rendered_count = 0
        self.fallback_count = 0
        self._known_pngs = set()  # 已识别的 PNG 文件，避免重复分配

    # ─── 工具检测 ───

    def _find_plantuml_cmd(self):
        """内置 plantuml.jar → Homebrew openjdk → 系统 java → 系统 plantuml"""
        bundled_jar = os.path.join(SCRIPTS_DIR, "plantuml.jar")
        java_cmd = None

        # 优先检测 Homebrew openjdk（macOS 上 /usr/bin/java 可能是 stub）
        homebrew_paths = [
            "/opt/homebrew/opt/openjdk/bin/java",       # Apple Silicon
            "/usr/local/opt/openjdk/bin/java",           # Intel Mac
        ]
        for path in homebrew_paths:
            if os.path.isfile(path):
                try:
                    result = subprocess.run(
                        [path, "-version"],
                        capture_output=True, text=True, timeout=5,
                    )
                    if result.returncode == 0:
                        java_cmd = path
                        break
                except (subprocess.TimeoutExpired, OSError):
                    pass

        # 如果 Homebrew openjdk 不可用，检测系统 java
        if not java_cmd:
            sys_java = shutil.which("java")
            if sys_java:
                try:
                    result = subprocess.run(
                        [sys_java, "-version"],
                        capture_output=True, text=True, timeout=5,
                    )
                    if result.returncode == 0:
                        java_cmd = sys_java
                except (subprocess.TimeoutExpired, OSError):
                    pass

        if os.path.isfile(bundled_jar) and java_cmd:
            return [java_cmd, "-jar", bundled_jar]
        if shutil.which("plantuml"):
            return ["plantuml"]

        # 提示安装 Java
        if not java_cmd and os.path.isfile(bundled_jar):
            print("[提示] PlantUML 渲染需要 Java，检测到 /usr/bin/java 为 macOS stub（不可运行）")
            print("  推荐安装: brew install openjdk")
        return None

    def _find_dot_cmd(self):
        """内置 graphviz → 系统 dot，返回 [cmd, renderer_flag]"""
        plat = platform.system().lower()
        if plat == "windows":
            bundled_dot = os.path.join(SCRIPTS_DIR, "graphviz", "win64", "bin", "dot.exe")
        else:
            bundled_dot = os.path.join(SCRIPTS_DIR, "graphviz", plat, "bin", "dot")
        # 平台专用渲染器：macOS 用 quartz（Core Text 支持 CJK），Linux 用 cairo（Pango 支持 CJK）
        renderer_flag = {
            "darwin": "-Tpng:quartz",
            "linux": "-Tpng:cairo",
            "windows": "-Tpng",
        }.get(plat, "-Tpng")
        if os.path.isfile(bundled_dot):
            return [bundled_dot, renderer_flag]
        if shutil.which("dot"):
            return ["dot", renderer_flag]
        return None

    # ─── CJK 字体检测 ───

    def _detect_cjk_font(self):
        """检测系统可用的 CJK 字体，返回字体名"""
        plat = platform.system()
        # 平台默认 CJK 字体优先级
        if plat == "Windows":
            candidates = ["Microsoft YaHei", "SimHei", "SimSun"]
        elif plat == "Darwin":
            candidates = ["PingFang SC", "Heiti SC", "STHeiti"]
        else:
            candidates = ["Noto Sans CJK SC", "WenQuanYi Micro Hei", "Droid Sans Fallback"]

        # 通过 fc-list（Linux/Mac）或注册表（Windows）验证字体可用性
        if plat == "Linux" or plat == "Darwin":
            try:
                result = subprocess.run(
                    ["fc-list", ":lang=zh", "family"],
                    capture_output=True, text=True, timeout=5,
                )
                available = result.stdout
                for font in candidates:
                    if font.lower() in available.lower():
                        return font
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
        elif plat == "Windows":
            # Windows 上常见 CJK 字体一般都可用，直接返回首选
            return candidates[0]

        # 回退：返回首选字体（渲染时若不可用则显示为方块，但比不设字体强）
        return candidates[0]

    # ─── 渲染入口 ───

    def render_code_block(self, lang, code):
        """渲染代码块为 PNG，返回 PNG 路径或 None"""
        import hashlib
        hash_str = hashlib.md5(code.encode("utf-8")).hexdigest()[:8]
        basename = lang + "_" + hash_str

        if lang in ("plantuml", "puml"):
            return self._render_plantuml(code, basename)
        elif lang in ("dot", "graphviz", "gv"):
            return self._render_graphviz(code, basename)
        elif lang == "mermaid":
            return self._render_mermaid(code, basename)
        else:
            return None

    def _render_plantuml(self, code, basename):
        """PlantUML → PNG（根据 @startuml 标识名定位输出文件）"""
        puml_path = os.path.join(self.diagrams_dir, basename + ".puml")

        # 注入 CJK 字体配置
        cjk_font = self._detect_cjk_font()
        font_line = "skinparam defaultFontName " + cjk_font + "\n"
        if "@startuml" in code and "skinparam defaultFontName" not in code:
            # 在 @startuml 后插入字体配置
            code = code.replace("@startuml", "@startuml\n" + font_line, 1)

        # 提取 @startuml 后的标识名，PlantUML 用此名命名输出 PNG
        name_match = re.search(r"@startuml\s+(\w+)", code)
        uml_name = name_match.group(1) if name_match else basename
        expected_png = os.path.join(self.diagrams_dir, uml_name + ".png")

        with open(puml_path, "w", encoding="utf-8") as f:
            f.write(code)

        cmd = self._find_plantuml_cmd()
        if cmd:
            # PlantUML 对相对路径会嵌套目录，必须使用绝对路径
            abs_puml_path = os.path.abspath(puml_path)
            abs_diagrams_dir = os.path.abspath(self.diagrams_dir)
            result = subprocess.run(
                cmd + ["-tpng", "-o", abs_diagrams_dir, abs_puml_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and os.path.isfile(expected_png) and os.path.getsize(expected_png) > 100:
                self.rendered_count += 1
                return expected_png
            # 如果期望文件名不匹配，扫描所有可能的新 PNG
            # PlantUML 有时用不同的命名规则（如含特殊字符时）
            for png_path in sorted(glob.glob(os.path.join(self.diagrams_dir, "*.png"))):
                if os.path.getsize(png_path) > 100 and png_path not in self._known_pngs:
                    self._known_pngs.add(png_path)
                    self.rendered_count += 1
                    return png_path

        self.fallback_count += 1
        return None

    def _render_graphviz(self, code, basename):
        """GraphViz DOT → PNG（平台专用渲染器确保 CJK 正确）"""
        dot_path = os.path.join(self.diagrams_dir, basename + ".dot")
        png_path = os.path.join(self.diagrams_dir, basename + ".png")

        # 注入 CJK 字体配置：替换不支持中文的字体名，或添加默认字体
        cjk_font = self._detect_cjk_font()
        # 替换已知不支持中文的字体
        for bad_font in ["Helvetica", "Arial", "DejaVu Sans", "Liberation Sans"]:
            code = code.replace('fontname="' + bad_font + '"', 'fontname="' + cjk_font + '"')
        # 如果全局没有 fontname，在首个 digraph/graph 行后注入 graph+node+edge 字体
        if "fontname=" not in code:
            code = re.sub(
                r"(digraph\s+\w+\s*\{)",
                r'\1\n    graph [fontname="' + cjk_font + '"]; node [fontname="' + cjk_font + '"]; edge [fontname="' + cjk_font + '"];',
                code, count=1,
            )

        with open(dot_path, "w", encoding="utf-8") as f:
            f.write(code)

        cmd = self._find_dot_cmd()
        if cmd:
            # cmd = [dot_binary, renderer_flag]，渲染器标志如 "-Tpng:quartz"
            result = subprocess.run(
                cmd + [dot_path, "-o", png_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and os.path.isfile(png_path) and os.path.getsize(png_path) > 100:
                self.rendered_count += 1
                return png_path
            if os.path.isfile(png_path) and os.path.getsize(png_path) <= 100:
                os.remove(png_path)

        self.fallback_count += 1
        return None

    def _render_mermaid(self, code, basename):
        """Mermaid → PNG（需 mmdc）"""
        png_path = os.path.join(self.diagrams_dir, basename + ".png")

        # 尝试系统 mmdc (mermaid-cli)
        mmdc_cmd = shutil.which("mmdc")
        if mmdc_cmd:
            mmd_path = os.path.join(self.diagrams_dir, basename + ".mmd")
            cjk_font = self._detect_cjk_font()
            # 生成 mmdc 配置 JSON（注入 CJK 字体）
            config_path = os.path.join(self.diagrams_dir, basename + "_mmdc_config.json")
            config_json = '{"theme": "default", "fontFamily": "' + cjk_font + '"}'
            with open(config_path, "w", encoding="utf-8") as f:
                f.write(config_json)
            with open(mmd_path, "w", encoding="utf-8") as f:
                f.write(code)
            result = subprocess.run(
                [mmdc_cmd, "-i", mmd_path, "-o", png_path, "-c", config_path],
                capture_output=True, text=True, timeout=30,
            )
            # 清理临时文件
            if os.path.isfile(mmd_path):
                os.remove(mmd_path)
            if os.path.isfile(config_path):
                os.remove(config_path)
            if result.returncode == 0 and os.path.isfile(png_path) and os.path.getsize(png_path) > 100:
                self.rendered_count += 1
                return png_path
            if os.path.isfile(png_path) and os.path.getsize(png_path) <= 100:
                os.remove(png_path)

        # Mermaid 渲染不可用，提示安装
        print("[提示] Mermaid 渲染需要 mmdc (mermaid-cli)，当前不可用")
        print("  推荐安装: npm install -g @mermaid-js/mermaid-cli")
        print("  或使用 PlantUML 替代（系统上下文→PlantUML Component，流程图→PlantUML Activity）")
        self.fallback_count += 1
        return None

    def summary(self):
        """输出渲染统计"""
        return self.rendered_count, self.fallback_count


# ─── Markdown → Docx 转换器 ───

class MarkdownToDocx:
    """Markdown → Docx 转换器，图表自动渲染嵌入"""

    def __init__(self, md_path, template_path=None, renderer=None):
        self.md_path = md_path
        self.md_dir = os.path.dirname(md_path) or "."
        self.renderer = renderer
        # 检测平台 CJK 字体，用于 Docx 文档样式
        self.cjk_font = DiagramRenderer._detect_cjk_font(self) if renderer else "Microsoft YaHei"
        # 如果没有 renderer，独立检测
        if not renderer:
            temp_renderer = DiagramRenderer(self.md_dir)
            self.cjk_font = temp_renderer._detect_cjk_font()
        if template_path and os.path.isfile(template_path):
            self.doc = Document(template_path)
            print("使用 Docx 模版: " + template_path)
        else:
            self.doc = Document()
            print("使用默认模版")
        print("Docx CJK字体: " + self.cjk_font)
        self._setup_styles()

    def _setup_styles(self):
        """配置 Docx 样式，使用检测到的 CJK 字体"""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.cjk_font
        font.size = Pt(10.5)
        rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
        if rFonts is not None:
            rFonts.set(qn("w:eastAsia"), self.cjk_font)

        for i in range(1, 5):
            heading_style = self.doc.styles["Heading " + str(i)]
            heading_font = heading_style.font
            heading_font.name = self.cjk_font
            if heading_style.element.rPr is not None:
                rFonts_h = heading_style.element.rPr.rFonts
                if rFonts_h is not None:
                    rFonts_h.set(qn("w:eastAsia"), self.cjk_font)

    def convert(self):
        """执行转换"""
        with open(self.md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        lines = [line.rstrip("\n") for line in lines]

        parser = _LineParser(lines)
        blocks = parser.parse()

        for block in blocks:
            self._render_block(block)

    def save(self, output_path):
        """保存 Docx"""
        self.doc.save(output_path)
        print("已保存: " + output_path + " (" + str(os.path.getsize(output_path)) + " bytes)")

    def _render_block(self, block):
        """渲染块级元素，图表代码块尝试渲染为 PNG 嵌入"""
        btype = block["type"]

        if btype == "heading":
            heading = self.doc.add_heading(level=min(block["level"], 4))
            self._add_inline_runs(heading, block["text"])

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                return
            ncols = len(rows[0])
            table = self.doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    cell = table.cell(i, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    self._add_inline_runs(p, cell_text)
                    if i == 0:
                        for run in p.runs:
                            run.bold = True

        elif btype == "code_block":
            lang = block.get("lang", "")
            code = block["code"]
            # 尝试渲染图表代码块为 PNG
            png_path = None
            if self.renderer and lang in ("plantuml", "puml", "dot", "graphviz", "gv", "mermaid"):
                png_path = self.renderer.render_code_block(lang, code)

            if png_path:
                # 渲染成功 → 嵌入 PNG 图片
                self._add_rendered_diagram(lang, png_path, code)
            else:
                # 渲染失败或非图表 → 保留源码文本
                self._add_code_paragraph(lang, code)

        elif btype == "callout":
            callout_type = block.get("callout_type", "NOTE")
            text = block["text"]
            p = self.doc.add_paragraph()
            prefix = "[" + callout_type + "] "
            run_prefix = p.add_run(prefix)
            run_prefix.bold = True
            if callout_type == "CAUTION":
                run_prefix.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)
            elif callout_type == "IMPORTANT":
                run_prefix.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            else:
                run_prefix.font.color.rgb = RGBColor(0x53, 0x34, 0x83)
            self._add_inline_runs(p, text)

        elif btype == "list":
            items = block["items"]
            ordered = block.get("ordered", False)
            for idx, item_text in enumerate(items):
                p = self.doc.add_paragraph()
                p.style = self.doc.styles["List Bullet"] if not ordered else self.doc.styles["List Number"]
                self._add_inline_runs(p, item_text)

        elif btype == "blockquote":
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            self._add_inline_runs(p, block["text"])

        elif btype == "image":
            self._add_image(block.get("alt", ""), block.get("path", ""))

        elif btype == "hr":
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif btype == "paragraph":
            text = block["text"]
            if text.strip():
                p = self.doc.add_paragraph()
                self._add_inline_runs(p, text)

    def _add_rendered_diagram(self, lang, png_path, original_code):
        """嵌入渲染后的 PNG 图片，下方附带源码摘要"""
        # 嵌入 PNG 图片
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(png_path, width=Inches(5))
        except Exception:
            p.text = ""
            alt_run = p.add_run("[图表渲染失败: " + lang + "]")
            alt_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            # 回退到文本
            self._add_code_paragraph(lang, original_code)
            return

        # 图片下方添加图表类型标签
        label_p = self.doc.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_p.add_run("[图表类型: " + lang + "]")
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _add_inline_runs(self, paragraph, text):
        """解析行内格式并添加 runs"""
        pattern = re.compile(
            r"(\*\*(.+?)\*\*)"       # bold
            r"|(\*(.+?)\*)"          # italic
            r"|(`(.+?)`)"            # inline code
            r"|(\[([^\]]+)\]\(([^)]+)\))"  # link
        )
        last_end = 0
        for m in pattern.finditer(text):
            if m.start() > last_end:
                paragraph.add_run(text[last_end:m.start()])
            if m.group(1):
                run = paragraph.add_run(m.group(2))
                run.bold = True
            elif m.group(3):
                run = paragraph.add_run(m.group(4))
                run.italic = True
            elif m.group(5):
                run = paragraph.add_run(m.group(6))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif m.group(7):
                display, url = m.group(8), m.group(9)
                run = paragraph.add_run(display)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
                run2 = paragraph.add_run(" (" + url + ")")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            last_end = m.end()
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_code_paragraph(self, lang, code):
        """添加代码块为文本段落"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if lang:
            label_run = p.add_run("[" + lang + "] ")
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        code_run = p.add_run(code)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(9)

    def _add_image(self, alt_text, image_path):
        """嵌入 PNG 图片"""
        candidates = [
            image_path,
            os.path.join(self.md_dir, image_path),
            os.path.join(self.md_dir, "diagrams", os.path.basename(image_path)),
        ]
        found_path = None
        for candidate in candidates:
            if os.path.isfile(candidate):
                found_path = candidate
                break
        if found_path:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(found_path, width=Inches(5))
            except Exception:
                p.text = ""
                alt_run = p.add_run("[图片: " + alt_text + " — " + found_path + "]")
                alt_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            p = self.doc.add_paragraph()
            alt_run = p.add_run("[图片未找到: " + image_path + "]")
            alt_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ─── Markdown 行解析器 ───

class _LineParser:
    """逐行解析 Markdown，产出块级元素列表"""

    def __init__(self, lines):
        self.lines = lines
        self.blocks = []
        self.pos = 0

    def parse(self):
        while self.pos < len(self.lines):
            line = self.lines[self.pos]

            if not line.strip():
                self.pos += 1
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                self.blocks.append({"type": "heading", "level": len(heading_match.group(1)), "text": heading_match.group(2)})
                self.pos += 1
                continue

            if re.match(r"^-{3,}$", line.strip()) or re.match(r"^={3,}$", line.strip()):
                self.blocks.append({"type": "hr"})
                self.pos += 1
                continue

            code_match = re.match(r"^```(\w*)$", line.strip())
            if code_match:
                lang = code_match.group(1)
                code_lines = []
                self.pos += 1
                while self.pos < len(self.lines):
                    if self.lines[self.pos].strip() == "```":
                        self.pos += 1
                        break
                    code_lines.append(self.lines[self.pos])
                    self.pos += 1
                self.blocks.append({"type": "code_block", "lang": lang, "code": "\n".join(code_lines)})
                continue

            if re.match(r"^\|", line):
                self._parse_table()
                continue

            callout_match = re.match(r"^>\s*\[!(NOTE|CAUTION|IMPORTANT|WARNING|TIP)\]\s*(.*)", line)
            if callout_match:
                callout_type, callout_text = callout_match.group(1), callout_match.group(2)
                extra_lines = []
                self.pos += 1
                while self.pos < len(self.lines):
                    next_line = self.lines[self.pos]
                    if next_line.startswith(">") and not re.match(r"^>\s*\[!", next_line):
                        extra_lines.append(next_line.lstrip("> ").strip())
                        self.pos += 1
                    else:
                        break
                self.blocks.append({"type": "callout", "callout_type": callout_type, "text": callout_text + ("\n" + "\n".join(extra_lines) if extra_lines else "")})
                continue

            if re.match(r"^>\s+", line):
                bq_lines = []
                while self.pos < len(self.lines) and self.lines[self.pos].startswith(">"):
                    bq_lines.append(self.lines[self.pos].lstrip("> ").strip())
                    self.pos += 1
                self.blocks.append({"type": "blockquote", "text": " ".join(bq_lines)})
                continue

            img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
            if img_match:
                self.blocks.append({"type": "image", "alt": img_match.group(1), "path": img_match.group(2)})
                self.pos += 1
                continue

            list_match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
            if list_match:
                self._parse_list(ordered=False)
                continue
            list_match_ordered = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
            if list_match_ordered:
                self._parse_list(ordered=True)
                continue

            self._parse_paragraph()

        return self.blocks

    def _parse_table(self):
        rows = []
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            if not re.match(r"^\|", line):
                break
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                self.pos += 1
                continue
            cells = [c.strip() for c in line.strip().split("|")]
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            rows.append(cells)
            self.pos += 1
        self.blocks.append({"type": "table", "rows": rows})

    def _parse_list(self, ordered=False):
        items = []
        bullet_re = r"^(\s*)([-*+])\s+(.+)$" if not ordered else r"^(\s*)(\d+)\.\s+(.+)$"
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            m = re.match(bullet_re, line)
            if m:
                items.append(m.group(3))
                self.pos += 1
            elif line.strip() == "":
                break
            else:
                break
        self.blocks.append({"type": "list", "ordered": ordered, "items": items})

    def _parse_paragraph(self):
        para_lines = []
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            if not line.strip():
                break
            if re.match(r"^(#{1,6})\s+", line):
                break
            if re.match(r"^```", line.strip()):
                break
            if re.match(r"^\|", line):
                break
            if re.match(r"^>\s+", line):
                break
            if re.match(r"^[-*+]\s+", line):
                break
            if re.match(r"^\d+\.\s+", line):
                break
            if re.match(r"^!\[", line.strip()):
                break
            para_lines.append(line)
            self.pos += 1
        if para_lines:
            self.blocks.append({"type": "paragraph", "text": " ".join(para_lines)})


# ─── 批量转换 ───

def convert_directory(output_dir, template=None):
    """批量转换 Markdown → Docx，图表自动渲染嵌入"""
    print("=== 转换 Markdown → Docx（图表自动渲染嵌入） ===")

    patterns = [
        os.path.join(output_dir, "requirement_analysis_spec_*.md"),
        os.path.join(output_dir, "functional_design_spec_*.md"),
    ]

    renderer = DiagramRenderer(output_dir)
    converted = 0

    for pattern in patterns:
        for md_file in sorted(glob.glob(pattern)):
            docx_file = os.path.splitext(md_file)[0] + ".docx"
            print("转换: " + md_file + " -> " + docx_file)
            try:
                converter = MarkdownToDocx(md_file, template, renderer)
                converter.convert()
                converter.save(docx_file)
                converted += 1
            except Exception as e:
                print("[错误] 转换失败: " + md_file)
                print("  " + str(e))

    rendered, fallback = renderer.summary()
    print("")
    print("=== 转换摘要 ===")
    print("转换文件数: " + str(converted))
    print("图表渲染: " + str(rendered) + " 个 PNG 嵌入, " + str(fallback) + " 个保留源码文本")
    print("Docx 文件目录: " + output_dir)

    docx_files = glob.glob(os.path.join(output_dir, "*.docx"))
    if docx_files:
        for f in sorted(docx_files):
            print("  " + os.path.basename(f) + " (" + str(os.path.getsize(f)) + " bytes)")
    else:
        print("(暂无 Docx 文件)")

    print("")
    print("图表渲染策略: 内置工具 → 系统工具 → 保留源码文本")
    print("Markdown 中图表源码始终保留，Docx 中成功渲染的图表以 PNG 展示")

    return converted


def convert_single(md_path, template=None):
    """转换单个 Markdown 文件 → Docx，图表自动渲染嵌入"""
    output_dir = os.path.dirname(md_path) or "."
    renderer = DiagramRenderer(output_dir)

    docx_file = os.path.splitext(md_path)[0] + ".docx"
    print("转换: " + md_path + " -> " + docx_file)
    converter = MarkdownToDocx(md_path, template, renderer)
    converter.convert()
    converter.save(docx_file)

    rendered, fallback = renderer.summary()
    print("图表渲染: " + str(rendered) + " 个 PNG 嵌入, " + str(fallback) + " 个保留源码文本")
    return docx_file


# ─── Main ───

if __name__ == "__main__":
    if len(sys.argv) >= 2 and os.path.isfile(sys.argv[1]) and sys.argv[1].endswith(".md"):
        # 单文件模式: python3 convert_docx.py <file.md> [template]
        convert_single(sys.argv[1], sys.argv[2] if len(sys.argv) >= 3 else None)
    else:
        # 目录批量模式: python3 convert_docx.py <dir> [template]
        output_dir = sys.argv[1] if len(sys.argv) >= 2 else "output"
        template = sys.argv[2] if len(sys.argv) >= 3 else None
        convert_directory(output_dir, template)